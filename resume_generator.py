
'-----------------------------------------------------------------------------'
import streamlit as st
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE
from datetime import datetime

# Function to add a bottom border to a cell
def add_bottom_border_to_cell(cell):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    bottomBorder = OxmlElement('w:bottom')
    bottomBorder.set(qn('w:val'), 'single')  # Change 'single' to other styles if needed
    bottomBorder.set(qn('w:sz'), '4')  # Border size (in eighths of a point)
    bottomBorder.set(qn('w:space'), '0')  # No space between borders and content
    bottomBorder.set(qn('w:color'), 'auto')  # Auto color or specify hex color
    tcBorders.append(bottomBorder)
    tcPr.append(tcBorders)

def set_vertical_alignment(cell, align="bottom"):
    tcPr = cell._element.get_or_add_tcPr()
    valign = OxmlElement('w:vAlign')
    valign.set(qn('w:val'), align)
    tcPr.append(valign)

# Function to set a cell's borders to None
def remove_borders_from_cell(cell, keep_bottom=False):
    # Create a new property element for the cell
    tcPr = cell._element.get_or_add_tcPr()

    # Specify the border elements to be removed (nil means no border)
    borders = ['top', 'left', 'right', 'insideH', 'insideV']
    if not keep_bottom:  # If not keeping the bottom border, add it to the list to be removed
        borders.append('bottom')

    for border in borders:
        tag = 'w:' + border
        element = OxmlElement(tag)
        element.set(qn('w:val'), 'nil')
        element.set(qn('w:sz'), '0')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), 'auto')
        tcPr.append(element)


def generate_resume(name, city, area_name, zipcode, email, phone, linkedin, summary,
                    programming_languages, libraries, business_intelligence, data_engineering,
                    big_data, profile, company_name, start_date, end_date, is_current_job, jd,
                    degree, university, certifications, additional_skills, statistical_methods,
                    data_collection, database_management, cloud_platforms, machine_learning):
    """Generates an ATS-friendly resume with a visually appealing layout."""
    doc = Document()

    # Set narrow margins for a sleek design
    for section in doc.sections:
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)

    # Set a professional font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Helper function to add and style cells
    def add_and_style_cell(text, bold=False, centered=False, font_size=12, color=None,is_heading=False,align_bottom_left=False,is_border=False):
        row = table.add_row()
        cell = row.cells[0]  # Add a new row and get the first cell
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = RGBColor(31, 78, 121)  # Dark blue color for headers
        if centered:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if is_heading:
            # Set a specific height for heading rows
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY  # Set the height rule
            row.height = Mm(10)  # Example height, adjust as needed
        if align_bottom_left:
            set_vertical_alignment(cell, "bottom")
        if is_border:
            add_bottom_border_to_cell(cell)

    # Initially, create a table with a single row for the header
    table = doc.add_table(rows=1,cols=1)
    table.style = 'Table Grid'

    # Add resume content
    add_and_style_cell(name, bold=True, centered=True, font_size=18, color=True)
    add_and_style_cell('Data Analyst', centered=True, color=True)
    contact_info = f"{city}, {area_name}, {zipcode} | {email} | {phone} | {linkedin}"
    add_and_style_cell(contact_info, centered=True)
    add_and_style_cell('SUMMARY', bold=True, color=True,font_size=12,is_heading=True,align_bottom_left=True,is_border=True)
    add_and_style_cell(summary,font_size=9)

    # Dynamic content from user input (skills, experience, etc.)
    add_and_style_cell('TECHNICAL SKILLS', bold=True, color=True,is_heading=True,align_bottom_left=True)
    skill_lines = [f"Programming Languages: {', '.join(programming_languages)}",
                   f"Libraries: {', '.join(libraries)}",
                   f"Business Intelligence: {', '.join(business_intelligence)}",
                   f"Data Engineering: {', '.join(data_engineering)}",
                   f"Big Data Tools: {', '.join(big_data)}",
                   f"Statistical Methods: {', '.join(statistical_methods)}",
                   f"Data Collection Techniques: {', '.join(data_collection)}",
                   f"Database Management Systems: {', '.join(database_management)}",
                   f"Cloud Platforms: {', '.join(cloud_platforms)}",
                   f"Machine Learning Libraries: {', '.join(machine_learning)}"]
    skill_lines = [f'• {line}' for line in skill_lines]
    # skill_lines = [line for line in skill_lines if not line.endswith(': ')]  # Filter out empty lines
    skills_text = '\n'.join(skill_lines)
    add_and_style_cell(skills_text,font_size=9)

    add_and_style_cell('PROFESSIONAL EXPERIENCE', bold=True, color=True,font_size=12,is_heading=True,align_bottom_left=True)
    experience_text = f"{profile} at {company_name} ({start_date.strftime('%B %Y')} - {'Present' if is_current_job else end_date.strftime('%B %Y')})"
    add_and_style_cell(experience_text)
    job_desc = '\n'.join([f'• {j}' for j in jd.split('\n')])
    add_and_style_cell(job_desc,font_size=9)  # Assuming 'jd' contains the job description

    add_and_style_cell('EDUCATION', bold=True, color=True,font_size=12,is_heading=True,align_bottom_left=True)
    education_text = f"{degree} from {university}"
    add_and_style_cell(education_text,font_size=9)

    # Handle certifications and additional skills similarly
    add_and_style_cell('CERTIFICATIONS', bold=True, color=True,font_size=12,is_heading=True,align_bottom_left=True)
    certs = certifications.split('\n')
    cer = '\n'.join([f'• {cert}' for cert in certs])
    add_and_style_cell(cer,font_size=9)

    add_and_style_cell('Additional Skills', bold=True, color=True,is_heading=True,align_bottom_left=True)
    skills = additional_skills.split('\n')
    skil = '\n'.join([f'• {skill}' for skill in skills])
    add_and_style_cell(skil,font_size=9)

    # Apply the no border function to each cell in the table

    # Assuming you have a list of section headings for which you want to keep the bottom border
    section_headings = ["SUMMARY", "TECHNICAL SKILLS", "PROFESSIONAL EXPERIENCE", "EDUCATION", "CERTIFICATIONS",
                        "Additional Skills"]

    # Loop through each cell in the table and remove borders accordingly
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()  # Get the cell's text content, trimming any leading/trailing whitespace
            if cell_text in section_headings:
                # If the cell's text is one of the section headings, keep the bottom border
                remove_borders_from_cell(cell, keep_bottom=True)
            else:
                # Otherwise, remove all borders
                remove_borders_from_cell(cell)
    row = table.rows[0]._element
    row.getparent().remove(row)



    # Save the document
    file_path = os.path.join(os.getcwd(), 'resume.docx')
    doc.save(file_path)
    return file_path

def main():
    st.title('ATS Friendly Resume Generator')
    with st.form("resume_form"):

        # UI for input fields
        with st.expander('Personal Details:'):
            name = st.text_input('Name')
            city = st.text_input('City')
            area_name = st.text_input('Area')
            zipcode = st.text_input('Zipcode')
            email = st.text_input('Email')
            phone = st.text_input('Phone')
            linkedin = st.text_input('LinkedIn')
            summary = st.text_area('Summary', placeholder='Write a Brief Summary')
        with st.expander('Skills'):
            statistical_methods = st.multiselect(label='Statistical methods',
                                                 options=['Statistical Techniques', 'Descriptive Statistics',
                                                          'Inferential Statistics', 'Probability Distribution',
                                                          'Hypothesis Testing', 'Regression'])
            programming_languages = st.multiselect(label='Programming Languages', options=['Python', 'SQL'])
            data_collection = st.multiselect(label='Data Collection',
                                             options=['requests', 'bs4', 'BeautifulSoup', 'lxml', 'API', 'web scraping'])
            libraries = st.multiselect(label='Libraries',
                                       options=['Numpy', 'Pandas', 'Matplotlib', 'Seaborn', 'Plotly', 'OpenCV'])
            database_management = st.multiselect(label='Database Management', options=['MySQL', 'MS SQL Server'])
            business_intelligence = st.multiselect(label='Business Intelligence',
                                                   options=['Power BI', 'DAX', 'Power Query', 'Tableau', 'Zoho',
                                                            'Quicksight', 'Google Studio', 'Excel Reporting'])
            data_engineering = st.multiselect(label='Data Engineering',
                                              options=['ETL Tools', 'SSIS', 'Data Warehouse', 'Data Pipeline',
                                                       'Data Mining', 'Data Wrangling', 'Data Munging'])
            cloud_platforms = st.multiselect(label='Cloud Platform',
                                             options=['Microsoft Azure', 'Azure Data Factory', 'Azure Cloud Services',
                                                      'Azure SQL Database', 'AWS', 'GCP'])
            big_data = st.multiselect(label='Big Data Tools', options=['Apache Spark', 'PySpark', 'Databricks'])
            machine_learning = st.multiselect(label='Machine Learning', options=['Scikit-learn'])
        with st.expander('Professional Experience'):
            profile = st.text_input('Previous Profile')
            company_name = st.text_input('Company Name')
            # Input fields for start and end date, and a checkbox for the current job
            start_date = st.date_input("Start Date")
            is_current_job = st.checkbox("Currently Working Here")
            end_date = st.date_input("End Date")

            jd = st.text_area('Explain Job Description')
        with st.expander('Education'):
            degree = st.text_input('Education')
            university = st.text_input('University')
            certifications = st.text_area('Certifications')
            additional_skills = st.text_area('Additional Skills')

        submitted = st.form_submit_button("Generate Resume")

    if submitted:
        # Call the function to generate resume
        # Make sure to pass all collected information to the function
        file_path = generate_resume(name, city, area_name, zipcode, email, phone, linkedin,
                                    summary, programming_languages, libraries, business_intelligence,
                                    data_engineering, big_data, profile, company_name,
                                    start_date, end_date, is_current_job, jd, degree, university,
                                    certifications, additional_skills, statistical_methods,
                                    data_collection, database_management, cloud_platforms,
                                    machine_learning)

        st.session_state['resume_path'] = file_path  # Store the path in session state

    # Check if the resume has been generated and offer a download button
    if 'resume_path' in st.session_state:
        with open(st.session_state['resume_path'], 'rb') as file:
            st.download_button(label="Download Resume", data=file, file_name="resume.docx",
                               mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    main()
