import streamlit as st
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def generate_resume(name, email, phone, linkedin, summary, programming_languages, business_intelligence, data_engineering, other_platforms):
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)

    # Add a table for the entire resume
    table = doc.add_table(rows=8, cols=1)
    table.style = 'Table Grid'

    # Name
    cell = table.cell(0, 0)
    cell.text = name
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell.paragraphs[0].runs[0].font.bold = True

    # Contact information
    cell = table.cell(1, 0)
    cell.text = f" {email}    {phone}   {linkedin}"
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Summary
    cell = table.cell(2, 0)
    cell.text = 'Summary'
    cell.paragraphs[0].runs[0].font.bold = True
    cell = table.cell(3, 0)
    cell.text = summary

    # Skills
    cell = table.cell(4, 0)
    cell.text = 'Skills'
    cell.paragraphs[0].runs[0].font.bold = True
    cell = table.cell(5, 0)
    cell.text = f"Programming Languages:\n{programming_languages}\n\nBusiness Intelligence:\n{business_intelligence}\n\nData Engineering:\n{data_engineering}\n\nOther Platforms:\n{other_platforms}"

    # Specify the absolute path to save the document
    file_path = os.path.join(os.getcwd(), 'resume.docx')
    doc.save(file_path)
    return file_path

def main():
    st.title('ATS Friendly Resume Generator')

    # Input fields
    name = st.text_input('Name')
    email = st.text_input('Email')
    phone = st.text_input('Phone')
    linkedin = st.text_input('LinkedIn')
    summary = st.text_area('Summary')
    programming_languages = st.text_area('Programming Languages')
    business_intelligence = st.text_area('Business Intelligence')
    data_engineering = st.text_area('Data Engineering')
    other_platforms = st.text_area('Other Platforms')

    # Generate resume
    if st.button('Generate Resume'):
        if name and email and phone and linkedin and summary and programming_languages and business_intelligence and data_engineering and other_platforms:
            file_path = generate_resume(name, email, phone, linkedin, summary, programming_languages, business_intelligence, data_engineering, other_platforms)
            st.success('Resume generated successfully!')
            st.download_button(
                label="Download your resume",
                data=open(file_path, 'rb').read(),
                file_name='resume.docx',
                mime='application/octet-stream'
            )

if __name__ == '__main__':
    main()
